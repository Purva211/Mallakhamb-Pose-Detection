from docx import Document
from docx.shared import Pt
import os

doc = Document()

# Add Title
title = doc.add_heading('Mallakhamb AI Pose Detection System Flow', 0)
title.alignment = 1 # Center

# Add Intro
doc.add_paragraph(
    "Here is the architectural and operational flow diagram for your project, "
    "outlining the three main stages: Data Pipeline, Model Training, and the Web Application Interface."
)

doc.add_heading('1. Data Pipeline', level=1)
p1 = doc.add_paragraph()
p1.add_run('Raw Images (1248) ').bold = True
p1.add_run('→ Preprocessing & Augmentation → ')
p1.add_run('Balanced Dataset (4750 Images, 19 Poses) ').bold = True
p1.add_run('→ Feature Extraction (MediaPipe Pose) → Normalization & Scaling → ')
p1.add_run('pose_landmarks.csv').bold = True

doc.add_heading('2. Model Training', level=1)
p2 = doc.add_paragraph()
p2.add_run('pose_landmarks.csv ').bold = True
p2.add_run('→ Train/Test Split & 5-Fold CV → Train Models (Random Forest, XGBoost, KNN, SVM) → ')
p2.add_run('Winning Model (Random Forest 98.43%) ').bold = True
p2.add_run('→ best_pose_classifier.pkl')

doc.add_heading('3. Web Application Interface', level=1)
doc.add_paragraph('This section handles real-time inference, including video processing and pole detection.')
flow = [
    "1. User Input (Image / Live Webcam / Video File)",
    "2. MediaPipe Pose Detection",
    "3. Is Human Detected?",
    "   - No → Error: No Human Detected",
    "   - Yes → OpenCV Pole Detection (Canny Edge & Hough Lines)",
    "4. Normalize User Skeleton",
    "5. Predict Pose (Random Forest Classifier)",
    "6. Verify Confidence:",
    "   - Is Confidence > 70% OR (Confidence > 52% & Pole Detected)?",
    "   - No → Reject Prediction",
    "   - Yes → Retrieve Gold Standard Ideal Pose",
    "7. Calculate Accuracy Score (Cosine Similarity)",
    "8. Calculate 3D Joint Angles",
    "9. Joint Angle Evaluation (Diff > 25°?):",
    "   - Yes → Generate Corrective Feedback (e.g., Bend your elbow)",
    "   - No → Generate Positive Feedback (Perfect form!)",
    "10. Display Results UI (Pose Name, Pole Status, Donut Chart, Tips)",
    "11. Text-to-Speech Output"
]
for item in flow:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Mermaid Diagram Code', level=1)
doc.add_paragraph(
    "Since Microsoft Word does not natively render Mermaid diagrams, you can copy the code below "
    "and paste it into a Mermaid viewer like https://mermaid.live to see the visual flowchart."
)

mermaid_code = """graph TD
    subgraph Data Pipeline
        A["Raw Images (1248)"] --> B["Preprocessing & Augmentation"]
        B --> C["Balanced Dataset<br>(4750 Images, 19 Poses)"]
        C --> D["Feature Extraction<br>(MediaPipe Pose)"]
        D --> E["Normalization & Scaling"]
        E --> F["pose_landmarks.csv"]
    end

    subgraph Model Training
        F --> G["Train/Test Split<br>& 5-Fold CV"]
        G --> H["Train Models<br>(Random Forest, XGBoost, KNN, SVM)"]
        H --> I["Winning Model<br>(Random Forest 98.43%)"]
        I --> J["best_pose_classifier.pkl"]
    end

    subgraph Web Application Interface
        K["User Input<br>(Image / Live Webcam / Video File)"] --> L["MediaPipe Pose Detection"]
        L --> M{"Is Human<br>Detected?"}
        M -- No --> N["Error: No Human Detected"]
        
        M -- Yes --> Pole["Pole Detection<br>(OpenCV Hough Lines)"]
        Pole --> O["Normalize User Skeleton"]
        
        O --> P["Predict Pose<br>(Random Forest)"]
        P --> Q{"Confidence > 70%<br>OR<br>(Confidence > 52% & Pole Detected)?"}
        Q -- No --> R["Reject Prediction"]
        Q -- Yes --> S["Retrieve Gold Standard<br>Ideal Pose"]
        
        S --> T["Calculate Accuracy<br>(Cosine Similarity)"]
        S --> U["Calculate 3D<br>Joint Angles"]
        
        U --> V{"Angle Diff > 25°?"}
        V -- Yes --> W["Generate Corrective Feedback<br>(e.g., Bend your elbow)"]
        V -- No --> X["Generate Positive Feedback<br>(Perfect form!)"]
        
        T --> Y["Display Results UI<br>(Pose Name, Pole Status, Donut Chart, Tips)"]
        W --> Y
        X --> Y
        Y --> Z["Text-to-Speech Output"]
    end

    J -.->|Loads Model| P
"""

code_para = doc.add_paragraph(mermaid_code)
code_para.style.font.name = 'Courier New'
code_para.style.font.size = Pt(9)

save_path = os.path.join(os.getcwd(), 'Project_Flow_Diagram.docx')
doc.save(save_path)
print(f"Document saved successfully at {save_path}")
